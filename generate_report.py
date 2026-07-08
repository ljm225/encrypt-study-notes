#!/usr/bin/env python3
"""生成登录密码加密传输漏洞实验报告 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页面边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def set_cell_shading(cell, color):
    """设置表格单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_heading_cn(text, level=1):
    """添加中文标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_para(text, bold=False, indent=False):
    """添加段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_code_block(text):
    """添加代码块样式"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('登录密码加密传输\n漏洞实验报告')
run.bold = True
run.font.size = Pt(26)
run.font.name = '黑体'
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('—— 基于 encrypt-study-notes 项目的安全实践')
run.font.size = Pt(14)
run.font.name = '宋体'
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for _ in range(6):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
today = datetime.date.today().strftime('%Y 年 %m 月 %d 日')
run = info.add_run(f'作者：刘佰鑫\n日期：{today}')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
add_heading_cn('目  录', level=1)
toc_items = [
    '一、实验概述',
    '    1.1 实验背景',
    '    1.2 实验目的',
    '    1.3 实验环境',
    '二、漏洞复现',
    '    2.1 漏洞描述',
    '    2.2 使用 Burp Suite 抓包演示',
    '    2.3 危害分析',
    '三、解决方案一：客户端哈希加密',
    '    3.1 原理',
    '    3.2 实现方式',
    '    3.3 安全性分析',
    '四、解决方案二：RSA 公钥加密',
    '    4.1 原理',
    '    4.2 实现方式',
    '    4.3 安全性分析',
    '五、解决方案三：AES-GCM 对称加密',
    '    5.1 原理',
    '    5.2 实现方式',
    '    5.3 安全性分析',
    '六、综合最佳实践方案',
    '    6.1 四层防护体系',
    '    6.2 攻击场景分析',
    '七、方案对比与总结',
    '    7.1 各方案对比',
    '    7.2 推荐建议',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# 一、实验概述
# ============================================================
add_heading_cn('一、实验概述', level=1)

add_heading_cn('1.1 实验背景', level=2)
add_para(
    '在 Web 应用的安全测试中，登录功能是最常见也是最重要的攻击面之一。'
    '许多 Web 应用在登录时直接将用户密码以明文方式通过 HTTP POST 请求传输，'
    '这使得攻击者可以通过网络抓包工具（如 Burp Suite、Wireshark 等）轻易获取用户密码。'
    '本实验针对这一安全漏洞，设计并实现了三种不同层级的加密传输方案，'
    '旨在探索如何在不依赖 HTTPS 的情况下（或在 HTTPS 之上额外增加安全层），'
    '保护用户密码在传输过程中的机密性。',
    indent=True
)

add_heading_cn('1.2 实验目的', level=2)
purposes = [
    '演示登录页面明文传输密码的安全漏洞及其危害；',
    '掌握客户端哈希加密（SHA-256 / PBKDF2）的实现与应用场景；',
    '掌握 RSA 公钥加密在密码传输中的应用；',
    '掌握 AES-GCM 对称加密的完整流程；',
    '理解各方案在不同攻击场景下的安全边界；',
    '形成 Web 应用密码传输的最佳实践方案。',
]
for p_text in purposes:
    add_para(f'• {p_text}')

add_heading_cn('1.3 实验环境', level=2)

# 环境表格
table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
env_data = [
    ('项目', '内容'),
    ('操作系统', 'Kali Linux'),
    ('Web 框架', '原生 HTML + JavaScript (Web Crypto API)'),
    ('抓包工具', 'Burp Suite Community Edition'),
    ('加密算法', 'SHA-256 / PBKDF2 / RSA-OAEP-2048 / AES-256-GCM'),
]
for i, (k, v) in enumerate(env_data):
    cell_k = table.cell(i, 0)
    cell_v = table.cell(i, 1)
    cell_k.text = k
    cell_v.text = v
    cell_k.paragraphs[0].runs[0].bold = True if i == 0 else False
    if i == 0:
        set_cell_shading(cell_k, '1A237E')
        set_cell_shading(cell_v, '1A237E')
        cell_k.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell_v.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_page_break()

# ============================================================
# 二、漏洞复现
# ============================================================
add_heading_cn('二、漏洞复现', level=1)

add_heading_cn('2.1 漏洞描述', level=2)
add_para(
    '未采取任何加密措施的登录页面，在用户提交表单时，密码字段以明文形式'
    '存在于 HTTP 请求体中。在 HTTP（非 HTTPS）协议下，整个通信过程是明文的，'
    '任何能够拦截网络流量的人都可以直接读取密码。即使在 HTTPS 环境下，'
    '如果攻击者通过社会工程学手段在用户设备上安装了自定义 CA 证书，'
    '同样可以通过中间人（MITM）攻击解密 TLS 流量，看到明文密码。',
    indent=True
)

add_heading_cn('2.2 使用 Burp Suite 抓包演示', level=2)
add_para('以下是一个典型的明文登录请求（使用 Burp Suite 拦截到的内容）：', indent=True)
add_code_block(
    'POST /login HTTP/1.1\n'
    'Host: example.com\n'
    'Content-Type: application/x-www-form-urlencoded\n\n'
    'username=admin&password=MySecretP@ss123'
)
add_para(
    '上例中，password=MySecretP@ss123 就是用户的明文密码，'
    '攻击者可以在 Burp Suite 的 Proxy → HTTP History 中直接看到。',
    indent=True
)

add_heading_cn('2.3 危害分析', level=2)
harms = [
    ('密码泄露', '攻击者直接获取用户密码，可登录受害者的账号'),
    ('撞库攻击', '多数用户在不同平台使用相同密码，一个泄露＝全线失守'),
    ('社会工程学', '攻击者可分析密码规律，推断用户个人信息'),
    ('合规风险', '违反《网络安全法》、《数据安全法》等法律法规对密码保护的要求'),
]
for h_title, h_desc in harms:
    add_para(f'• {h_title}：{h_desc}')

doc.add_page_break()

# ============================================================
# 三、客户端哈希加密
# ============================================================
add_heading_cn('三、解决方案一：客户端哈希加密', level=1)

add_heading_cn('3.1 原理', level=2)
add_para(
    '客户端哈希加密的原理是在浏览器端对用户密码进行哈希运算，'
    '然后将哈希值而非明文密码发送给服务器。即使攻击者截获了哈希值，'
    '也无法直接获得原始密码。本方案使用 Web Crypto API 提供的 '
    'SHA-256 和 PBKDF2（基于密码的密钥派生函数）两种算法。',
    indent=True
)

add_heading_cn('3.2 实现方式', level=2)
add_para('核心流程：', bold=True)
steps = [
    '步骤1：页面加载时，服务器下发一个随机盐值（Salt）',
    '步骤2：用户输入密码，点击登录',
    '步骤3：前端计算 hash = SHA-256(Salt + Password)',
    '步骤4：将 { username, password_hash } 发送给服务器',
    '步骤5：服务器使用相同的盐值和算法验证哈希值',
]
for s in steps:
    add_para(f'  {s}')

doc.add_paragraph()
add_para('关键代码示例（使用 PBKDF2，推荐方案）：', bold=True)
add_code_block(
    'async function hashWithPBKDF2(password, salt) {\n'
    '  const encoder = new TextEncoder();\n'
    '  const keyMaterial = await crypto.subtle.importKey(\n'
    '    "raw", encoder.encode(password), "PBKDF2",\n'
    '    false, ["deriveBits"]\n'
    '  );\n'
    '  const hashBuffer = await crypto.subtle.deriveBits(\n'
    '    { name: "PBKDF2", salt: encoder.encode(salt),\n'
    '      iterations: 100000, hash: "SHA-256" },\n'
    '    keyMaterial, 256\n'
    '  );\n'
    '  return Array.from(new Uint8Array(hashBuffer))\n'
    '    .map(b => b.toString(16).padStart(2, "0")).join("");\n'
    '}'
)

add_heading_cn('3.3 安全性分析', level=2)
add_para('✅ 优点：', bold=True)
for item in [
    '密码明文不会出现在网络请求中',
    'PBKDF2 10万次迭代，暴力破解成本极高',
    '加盐防止彩虹表攻击',
    '实现简单，性能消耗低',
]:
    add_para(f'  • {item}')

add_para('❌ 局限性：', bold=True)
for item in [
    '哈希值本身如果被重放攻击利用，等同于密码泄露',
    '需要服务器配合改为哈希验证逻辑',
    '无防重放机制（需额外加时间戳 / Nonce）',
]:
    add_para(f'  • {item}')

doc.add_page_break()

# ============================================================
# 四、RSA 公钥加密
# ============================================================
add_heading_cn('四、解决方案二：RSA 公钥加密', level=1)

add_heading_cn('4.1 原理', level=2)
add_para(
    'RSA 是一种非对称加密算法。服务器生成一对密钥（公钥 + 私钥），'
    '公钥公开发送给前端，私钥由服务器秘密保管。前端使用公钥加密密码，'
    '只有持有私钥的服务器才能解密。即使攻击者截获了加密后的密文，'
    '由于没有私钥，也无法还原出原始密码。本实现使用 RSA-OAEP 填充模式，'
    '密钥长度 2048 位，哈希算法 SHA-256。',
    indent=True
)

add_heading_cn('4.2 实现方式', level=2)
add_para('核心流程：', bold=True)
steps = [
    '步骤1：服务器生成 RSA-2048 密钥对，私钥安全存储',
    '步骤2：页面加载时，服务器将公钥下发给前端',
    '步骤3：用户输入密码，前端用公钥加密',
    '步骤4：将 { username, encrypted_password } 发送给服务器',
    '步骤5：服务器用私钥解密，得到明文密码后验证',
]
for s in steps:
    add_para(f'  {s}')

doc.add_paragraph()
add_para('关键代码示例：', bold=True)
add_code_block(
    '// 使用 RSA-OAEP 公钥加密密码\n'
    'async function encryptWithRSA(publicKey, plaintext) {\n'
    '  const encoder = new TextEncoder();\n'
    '  const data = encoder.encode(plaintext);\n'
    '  const encrypted = await crypto.subtle.encrypt(\n'
    '    { name: "RSA-OAEP" },\n'
    '    publicKey,\n'
    '    data\n'
    '  );\n'
    '  return arrayBufferToBase64(encrypted);\n'
    '}'
)

add_heading_cn('4.3 安全性分析', level=2)
add_para('✅ 优点：', bold=True)
add_para('  • 私钥在服务器端，截获密文也无法解密')
add_para('  • RSA-2048 在当前计算能力下不可破解')
add_para('  • 无需服务器存储共享密钥，公钥可随意分发')

add_para('❌ 局限性：', bold=True)
add_para('  • RSA 加密运算较慢，不适合加密大数据')
add_para('  • 必须配合 HTTPS 使用，否则公钥可能被替换（中间人攻击）')
add_para('  • 同样需要防重放机制')

doc.add_page_break()

# ============================================================
# 五、AES-GCM 对称加密
# ============================================================
add_heading_cn('五、解决方案三：AES-GCM 对称加密', level=1)

add_heading_cn('5.1 原理', level=2)
add_para(
    'AES-GCM（Galois/Counter Mode）是一种对称加密算法，'
    '使用同一个密钥进行加密和解密。与 RSA 不同，对称加密速度快、'
    '适合加密任意长度的数据。GCM 模式在加密的同时提供认证标签（Authentication Tag），'
    '可以检测密文是否被篡改，同时保证机密性和完整性。本实现使用 AES-256 密钥长度。',
    indent=True
)

add_heading_cn('5.2 实现方式', level=2)
add_para('核心流程：', bold=True)
steps = [
    '步骤1：服务器生成 AES-256 密钥，通过 HTTPS 安全下发给前端',
    '步骤2：前端生成 12 字节随机 IV（初始化向量）',
    '步骤3：使用 AES-GCM 加密密码，得到 { 密文, 认证标签 }',
    '步骤4：将 { ciphertext, iv, auth_tag } 发送给服务器',
    '步骤5：服务器使用密钥 + IV 解密，并验证认证标签',
]
for s in steps:
    add_para(f'  {s}')

doc.add_paragraph()
add_para('关键代码示例：', bold=True)
add_code_block(
    'async function encryptAESGCM(key, plaintext) {\n'
    '  const iv = crypto.getRandomValues(new Uint8Array(12));\n'
    '  const encrypted = await crypto.subtle.encrypt(\n'
    '    { name: "AES-GCM", iv: iv, tagLength: 128 },\n'
    '    key,\n'
    '    new TextEncoder().encode(plaintext)\n'
    '  );\n'
    '  // 输出: 密文 || 认证标签（最后16字节）\n'
    '  return { ciphertext, authTag, iv };\n'
    '}'
)

add_heading_cn('5.3 安全性分析', level=2)
add_para('✅ 优点：', bold=True)
add_para('  • AES-256-GCM 是目前最安全的对称加密模式之一')
add_para('  • 加密速度快，适合大规模应用')
add_para('  • 自带认证标签，防止密文被篡改')
add_para('  • 每次加密使用不同的随机 IV，密文不重复')

add_para('❌ 局限性：', bold=True)
add_para('  • 密钥分发依赖 HTTPS，无 HTTPS 则无安全性')
add_para('  • 密钥管理复杂（密钥分发、轮换、撤销）')
add_para('  • IV 不允许重复使用，否则加密失效')

doc.add_page_break()

# ============================================================
# 六、综合最佳实践方案
# ============================================================
add_heading_cn('六、综合最佳实践方案', level=1)

add_heading_cn('6.1 四层防护体系', level=2)
add_para(
    '在实际生产环境中，单一加密方案往往无法应对所有攻击场景。'
    '本实验最终设计了一个四层防护体系，从内到外层层加固：',
    indent=True
)

# 分层表格
table2 = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
layer_data = [
    ('层级', '技术方案', '防护目标'),
    ('第1层', 'PBKDF2 加盐哈希\n（10万次迭代）', '防止密码明文出现在\n网络请求中'),
    ('第2层', '时间戳 + 随机 Nonce', '防止重放攻击'),
    ('第3层', 'RSA-2048 公钥加密', '防止窃听（无私钥不解密）'),
    ('第4层', 'HTTPS TLS 1.3', '防止中间人攻击'),
]
for i, (a, b, c) in enumerate(layer_data):
    table2.cell(i, 0).text = a
    table2.cell(i, 1).text = b
    table2.cell(i, 2).text = c
    if i == 0:
        for j in range(3):
            set_cell_shading(table2.cell(i, j), '1A237E')
            table2.cell(i, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            table2.cell(i, j).paragraphs[0].runs[0].bold = True

add_heading_cn('6.2 攻击场景分析', level=2)

# 攻击场景表格
table3 = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
attack_data = [
    ('攻击方式', '能否破解'),
    ('Burp Suite 无证书 MITM', '❌ HTTPS 直接拦截'),
    ('Burp Suite + 伪造 CA 证书', '❌ 看到 RSA 密文，无法解密'),
    ('重放攻击（Replay）', '❌ 时间戳 ±5 分钟 + Nonce 过期'),
    ('彩虹表 / 暴力破解', '❌ PBKDF2 加盐 + 10万次迭代'),
]
for i, (a, b) in enumerate(attack_data):
    table3.cell(i, 0).text = a
    table3.cell(i, 1).text = b
    if i == 0:
        set_cell_shading(table3.cell(i, 0), '1A237E')
        set_cell_shading(table3.cell(i, 1), '1A237E')
        table3.cell(i, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        table3.cell(i, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        table3.cell(i, 0).paragraphs[0].runs[0].bold = True
        table3.cell(i, 1).paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ============================================================
# 七、方案对比与总结
# ============================================================
add_heading_cn('七、方案对比与总结', level=1)

add_heading_cn('7.1 各方案对比', level=2)

# 对比表格
table4 = doc.add_table(rows=5, cols=6, style='Light Grid Accent 1')
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
compare_data = [
    ('方案', '防窃听', '防重放', '服务端改动', '性能', '实现难度'),
    ('客户端哈希', '⭐⭐⭐', '❌ 需追加', '需改为哈希验证', '⚡ 快', '⭐ 简单'),
    ('RSA 加密', '⭐⭐⭐⭐', '❌ 需追加', '需解密操作', '🐢 较慢', '⭐⭐⭐'),
    ('AES-GCM', '⭐⭐⭐⭐', '❌ 需追加', '需解密操作', '⚡ 快', '⭐⭐'),
    ('四层综合', '⭐⭐⭐⭐⭐', '✅ 内置', '需全部改造', '适中', '⭐⭐⭐⭐'),
]
for i, row in enumerate(compare_data):
    for j, val in enumerate(row):
        table4.cell(i, j).text = val
        if i == 0:
            set_cell_shading(table4.cell(i, j), '1A237E')
            table4.cell(i, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            table4.cell(i, j).paragraphs[0].runs[0].bold = True

add_heading_cn('7.2 推荐建议', level=2)

recommendations = [
    ('基础防护（入门级）', '客户端哈希（SHA-256 + Salt）+ HTTPS。实现简单，适用于低风险场景。'),
    ('中等防护（推荐）', 'RSA 公钥加密 + HTTPS。安全性较高，适用于大多数 Web 应用。'),
    ('高级防护（最佳实践）', '四层综合方案：PBKDF2 → 时间戳/Nonce → RSA 加密 → HTTPS。适用于金融、政务等高安全需求场景。'),
    ('核心原则', '永远不要只依赖客户端加密。HTTPS 是所有加密方案的基础，客户端加密是在 HTTPS 之上的额外安全层。'),
]

for r_title, r_desc in recommendations:
    add_para(f'• {r_title}', bold=True)
    add_para(f'  {r_desc}')
    doc.add_paragraph()

# ============================================================
# 结语
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
add_heading_cn('结  语', level=1)
add_para(
    '本实验从实际安全测试中遇到的明文传输漏洞出发，'
    '系统地设计了三种不同层级的加密传输方案，并最终给出了一套综合最佳实践。'
    '实验成果已整理为开源项目 encrypt-study-notes，'
    '包含完整的可运行的 HTML 演示页面，可供后续学习和参考。'
    '在 Web 安全领域，没有一劳永逸的解决方案，只有持续加固、'
    '纵深防御才是保障用户数据安全的正道。',
    indent=True
)

doc.add_paragraph()
add_para(f'报告完成日期：{today}')
add_para('项目地址：https://github.com/ljm225/encrypt-study-notes')

# ============================================================
# 保存
# ============================================================
output_path = '/root/Desktop/登录密码加密传输漏洞实验报告.docx'
doc.save(output_path)
print(f'✅ 报告已生成：{output_path}')
